from docx import Document
import datetime

document = Document('word_data/範例.docx')
paragraphs = document.paragraphs

#Store content of second paragraph
text = paragraphs[0].text

#Clear content
paragraphs[0]._p.clear()

#Recreate second paragraph
paragraphs[0].add_run('吳先生' + text+"\n")
paragraphs[0].add_run("這是報價單")

tomorrow = datetime.date.today() + datetime.timedelta(days=10)

#add something in last value

p = document.add_paragraph('如果沒問題，請在'+str(tomorrow)+"回傳")
p = document.add_paragraph("經辦人: 吳如峰")

document.save('新文件.docx')

