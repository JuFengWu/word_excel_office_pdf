import docx

doc = docx.Document('word_data/測試.docx')

p = doc.add_paragraph('如果沒問題，請在'+"今天回傳")
p = doc.add_paragraph("經辦人: 吳如峰")

doc.save('word_data/新文件2.docx')