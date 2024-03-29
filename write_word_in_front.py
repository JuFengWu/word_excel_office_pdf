import docx

doc = docx.Document('word_data/測試.docx')
paragraphs = doc.paragraphs
paragraphs[0].add_run("吳先生\n")
paragraphs[0].add_run("這是報價單")

doc.save('word_data/新文件.docx')