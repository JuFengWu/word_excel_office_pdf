import docx

doc = docx.Document('word_data/測試.docx')
paragraphs = doc.paragraphs
r = paragraphs[0].add_run()
r.add_picture('logo.JPG')

doc.save('word_data/新文件4.docx')