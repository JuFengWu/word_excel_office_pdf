import docx

doc = docx.Document('word_data/測試.docx')

print('段落數量： ', len(doc.paragraphs))

for para in doc.paragraphs:
    print(para.text)

para = doc.paragraphs[2]
print("-------")
print("段落的內容為\n\n")
print(para.text + '\n')
print('run數量： ', len(para.runs))

for i in range(0, len(para.runs)):
    print(i, para.runs[i].text)
