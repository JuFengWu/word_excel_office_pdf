import docx

doc = docx.Document('word_data/測試.docx')

print('段落數量： ', len(doc.paragraphs))

for para in doc.paragraphs:
    print(para.text)
    print("--------")