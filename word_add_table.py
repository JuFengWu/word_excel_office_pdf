from docx import Document
from docx.enum.style import WD_STYLE_TYPE

doc = Document('word_data/測試.docx')

table = doc.add_table(rows=1, cols=4)
table.style = 'TGrid'

items = {"星星切磨寶石": 3400, "鳥圖形造型寶石": 4300 , "項鍊" : 2500}

tableTitle = table.rows[0].cells
tableTitle[0].text = '品項'
tableTitle[1].text = '金額'
tableTitle[2].text = '稅金'
tableTitle[3].text = '總金額'


for item in items:
    row_cells = table.add_row().cells
    row_cells[0].text = item
    row_cells[1].text = str(items[item])
    row_cells[2].text = str(int(items[item]) * 0.05)
    row_cells[3].text = str(int(items[item]) * 1.05)


p = doc.add_paragraph("經辦人: 吳如峰")
doc.save('word_data/新文件3.docx')

